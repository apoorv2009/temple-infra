from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = r"D:\Projects\MyTemple\temple-infra\docs\temple-ai-rag-agentic-ui-guide.docx"

ACCENT = RGBColor(230, 117, 39)
ACCENT_DARK = RGBColor(122, 64, 28)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(96, 96, 96)
LIGHT_FILL = "FDF4EE"
MID_FILL = "F7E1D4"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_page_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def style_run(run, *, size=11, bold=False, color=TEXT, font="Aptos"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.color.rgb = color


def add_paragraph(doc, text="", *, style=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        style_run(run)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = 1.15
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(12 if level == 1 else 8)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    style_run(
        run,
        size=18 if level == 1 else 13,
        bold=True,
        color=ACCENT_DARK if level == 1 else TEXT,
    )
    return p


def add_bullets(doc, items, *, indent=0.2):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        style_run(run)


def add_numbered(doc, items, *, indent=0.2):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        style_run(run)


def add_info_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    style_run(r1, size=11, bold=True, color=ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body)
    style_run(r2, size=10.5, color=TEXT)
    add_paragraph(doc, "", space_after=4)


def add_two_col_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], MID_FILL)
        set_cell_margins(hdr[i], top=120, start=110, bottom=110, end=110)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        style_run(r, size=10.5, bold=True, color=ACCENT_DARK)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_margins(cells[i], top=95, start=110, bottom=95, end=110)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            style_run(r, size=10.2)
    add_paragraph(doc, "", space_after=5)


def add_cover(doc):
    section = doc.sections[0]
    set_page_margins(section)

    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    set_cell_shading(cell, "FFF1E8")
    set_cell_margins(cell, top=240, start=220, bottom=220, end=220)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Temple App AI Architecture Guide")
    style_run(r, size=24, bold=True, color=ACCENT_DARK)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("RAG pipeline, agentic UI, service design, implementation steps, and learning material")
    style_run(r2, size=11.5, color=MUTED)

    add_paragraph(doc, "", space_after=12)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    entries = [
        ("Prepared for", "Temple App product and engineering team"),
        ("Primary target", "Temple-scoped AI assistant for devotees and temple admins"),
        ("Recommended stack", "OpenAI Responses API + GPT-4.1 + text-embedding-3-small + pgvector"),
    ]
    for row_idx, (left, right) in enumerate(entries):
        c1, c2 = meta.rows[row_idx].cells
        for c in (c1, c2):
            set_cell_margins(c, top=120, start=120, bottom=120, end=120)
        set_cell_shading(c1, MID_FILL)
        r1 = c1.paragraphs[0].add_run(left)
        style_run(r1, size=10.5, bold=True, color=ACCENT_DARK)
        r2 = c2.paragraphs[0].add_run(right)
        style_run(r2, size=10.5)

    add_paragraph(doc, "", space_after=16)
    add_info_box(
        doc,
        "Recommended starting point",
        "Build a new temple-ai-service. Keep retrieval and orchestration outside the existing booking, admin, registration, and identity services. This keeps AI isolated, easier to evolve, and simpler to secure.",
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)

    add_cover(doc)

    add_heading(doc, "1. Why AI is useful in the current app", 1)
    add_bullets(
        doc,
        [
            "It can answer temple-specific questions using your own data instead of generic internet answers.",
            "It can guide devotees through bookings, donations, and membership flows without making them search multiple screens.",
            "It can help admins draft notifications, summarize temple activity, and review member actions faster.",
            "It can become a temple-scoped assistant instead of a generic chatbot, which is where RAG matters most.",
        ],
    )
    add_info_box(
        doc,
        "What you should build first",
        "Start with a temple assistant in the Chat tab. Let it answer questions from temple notices, FAQs, timings, donation instructions, booking rules, and user-specific status APIs.",
    )

    add_heading(doc, "2. Recommended technology stack", 1)
    add_two_col_table(
        doc,
        [
            ("Mobile UI", "Existing React Native + Expo app. Add a dedicated Chat tab with card-based responses and action buttons."),
            ("AI orchestration", "New Python FastAPI microservice named temple-ai-service."),
            ("Generation model", "OpenAI GPT-4.1 for the first version. Good tool calling, broad context handling, and stable instruction following."),
            ("Embedding model", "text-embedding-3-small first. Low cost and strong enough for temple content retrieval."),
            ("Higher retrieval option", "text-embedding-3-large later if precision becomes a problem for multilingual or subtle semantic search."),
            ("Vector store", "Postgres with pgvector in a separate temple_ai database."),
            ("API style", "OpenAI Responses API for tool calling and agent loops."),
            ("Streaming", "Optional in phase 2. Start with non-streaming responses, then add streaming later if chat UX needs it."),
        ],
        ["Layer", "Recommendation"],
    )

    add_heading(doc, "3. Model strategy", 1)
    add_bullets(
        doc,
        [
            "Use embeddings to retrieve relevant data. Embeddings are not the answer model; they are the search representation.",
            "Use GPT-4.1 to read retrieved context, decide whether tools are needed, and generate the final answer.",
            "Use structured tool calls when the assistant must read live user data such as donation status, booking status, or membership status.",
        ],
    )
    add_two_col_table(
        doc,
        [
            ("What retrieves data?", "Embeddings + vector similarity search in pgvector."),
            ("What writes the answer?", "GPT-4.1 through the Responses API."),
            ("What checks live app state?", "Tool calls from temple-ai-service to existing microservices."),
            ("When to consider GPT-5 later?", "When you want more advanced multi-step reasoning, better planning, or more complex admin workflows."),
        ],
        ["Question", "Answer"],
    )

    add_heading(doc, "4. Target architecture", 1)
    add_bullets(
        doc,
        [
            "Frontend never talks directly to OpenAI.",
            "Frontend talks to temple-ai-service.",
            "temple-ai-service retrieves temple knowledge from pgvector and calls OpenAI.",
            "temple-ai-service also calls existing microservices when live state is needed.",
        ],
    )
    add_two_col_table(
        doc,
        [
            ("Temple App", "Collects user question, shows answers, shows action cards, and launches flows such as Book or Donate."),
            ("temple-ai-service", "Owns RAG pipeline, tool calling, prompt construction, response formatting, and conversation policy."),
            ("temple_ai database", "Stores source documents, chunks, metadata, embeddings, conversation references, and retrieval audit data."),
            ("temple-admin-service", "Provides temple notices, wall of fame, timings, policies, payment instructions, slot rules, and future admin tools."),
            ("temple-registration-service", "Provides booking status, donation status, membership status, and user transaction history."),
            ("temple-identity-service", "Provides user profile, role, approved primary temple, and security checks."),
            ("OpenAI API", "Provides embeddings and response generation with tool orchestration."),
        ],
        ["Component", "Role"],
    )

    add_heading(doc, "5. Where the RAG pipeline should be created", 1)
    add_paragraph(
        doc,
        "Create the RAG pipeline inside a new microservice named temple-ai-service. Do not mix it into the current gateway or booking services. Retrieval, prompt building, AI safety rules, and tool orchestration should stay isolated there.",
        space_after=8,
    )
    add_numbered(
        doc,
        [
            "Ingest temple content from existing services or admin uploads.",
            "Normalize the text and split it into chunks.",
            "Generate embeddings for each chunk.",
            "Store chunks and vectors with metadata such as temple_id, content_type, language, source_id, and visibility.",
            "At question time, embed the user query and search only within the selected temple scope.",
            "Assemble top results into the prompt.",
            "If the query needs live data, call tools before final response generation.",
            "Return a final answer plus UI actions such as Open Booking, Open Donate, or Draft Notification.",
        ],
    )

    add_heading(doc, "6. Detailed RAG pipeline", 1)
    add_two_col_table(
        doc,
        [
            ("Ingestion", "Read content from temple notices, wall of fame, FAQs, rules, temple profile, donation instructions, and booking instructions."),
            ("Chunking", "Break larger documents into small, meaningful passages. Target semantic chunks, not fixed arbitrary blocks only."),
            ("Embedding", "Use text-embedding-3-small to convert each chunk into a vector."),
            ("Storage", "Store raw chunk text, metadata, and vector in temple_ai Postgres with pgvector."),
            ("Retrieval", "On each user query, embed the query and retrieve top-k chunks filtered by temple_id and visibility."),
            ("Optional reranking", "Not needed in v1. Add later only if answer quality is weak."),
            ("Prompt assembly", "Build a system prompt, user query, retrieved context, and any live tool results into one final model request."),
            ("Answer generation", "Use GPT-4.1 to answer only from retrieved and tool-backed context."),
            ("Citations", "Return source snippets or source titles to reduce hallucination and improve trust."),
        ],
        ["Stage", "What it does"],
    )

    add_info_box(
        doc,
        "Important design rule",
        "Temple knowledge retrieval must always be filtered by temple_id. A devotee of one temple must not receive content from another temple unless you deliberately build a cross-temple mode later.",
    )

    add_heading(doc, "7. Agentic UI in this app", 1)
    add_paragraph(
        doc,
        "Agentic UI means the assistant does not only return text. It can also return structured cards, buttons, confirmations, and next actions. In your app, this is more useful than plain chatbot text.",
        space_after=8,
    )
    add_two_col_table(
        doc,
        [
            ("Devotee question", "What is the Shantidhara process?"),
            ("Assistant response", "Short answer + Book Shantidhara button + Open temple rules link."),
            ("Devotee question", "I want to donate 501 rupees."),
            ("Assistant response", "Donation summary card + Open Donate flow button."),
            ("Admin question", "Draft a Paryushan information notice."),
            ("Assistant response", "Draft notification card + Edit + Publish action."),
            ("Admin question", "Show pending payment proofs."),
            ("Assistant response", "Review queue card with open actions."),
        ],
        ["Example input", "Agentic UI output"],
    )

    add_heading(doc, "8. Best AI features to build in the current app", 1)
    add_bullets(
        doc,
        [
            "Temple Help Assistant for devotees: timings, booking rules, donation instructions, FAQs, notices.",
            "Personal Temple Assistant: membership status, booking status, donation proof status, what to do next.",
            "Admin Drafting Assistant: draft Information and Wall of Fame posts.",
            "Admin Operations Assistant: summarize pending requests, payment proofs, and recent notifications.",
            "Future multilingual assistant: Hindi, English, and local language support if temple communities require it.",
        ],
    )

    add_heading(doc, "9. Suggested database design for temple-ai-service", 1)
    add_two_col_table(
        doc,
        [
            ("source_documents", "One row per temple document or content source. Fields: document_id, temple_id, source_type, source_ref, title, raw_text, language, visibility, version, created_at, updated_at."),
            ("source_chunks", "Chunk-level rows. Fields: chunk_id, document_id, temple_id, chunk_text, chunk_index, token_estimate, embedding vector, metadata_json."),
            ("chat_sessions", "Session tracking. Fields: session_id, user_id, temple_id, role, created_at, updated_at."),
            ("chat_messages", "User and assistant messages. Fields: message_id, session_id, sender_type, content, tool_calls_json, citations_json, created_at."),
            ("retrieval_logs", "Audit table for which chunks were used. Helpful for debugging and quality review."),
            ("tool_audit_logs", "Audit of which service tools were called by the assistant and why."),
        ],
        ["Table", "Purpose"],
    )

    add_heading(doc, "10. Tool layer design", 1)
    add_paragraph(
        doc,
        "RAG should answer knowledge questions, but user-specific and real-time answers must come from tools. This prevents stale answers and reduces hallucination.",
        space_after=8,
    )
    add_two_col_table(
        doc,
        [
            ("get_membership_status(user_id, temple_id)", "registration service"),
            ("get_booking_status(user_id, temple_id)", "registration service"),
            ("get_donation_status(user_id, temple_id)", "registration service"),
            ("get_available_shantidhara_slots(temple_id, date)", "admin service"),
            ("get_temple_profile(temple_id)", "admin service"),
            ("draft_information_notice(context)", "AI only, no external tool needed"),
            ("publish_information_notice(temple_id, payload)", "admin service, admin-only"),
            ("publish_wall_of_fame(temple_id, payload)", "admin service, admin-only"),
        ],
        ["Tool", "Owned by"],
    )

    add_heading(doc, "11. Security and governance", 1)
    add_bullets(
        doc,
        [
            "Always enforce role checks before tool execution.",
            "Always enforce temple_id filtering on retrieval.",
            "Never let the model call publish tools without explicit admin permissions and confirmation UI.",
            "Do not include raw secrets, tokens, internal IDs, or payment screenshots in retrieval context.",
            "Keep an audit trail of tool calls and retrieved chunks for production support.",
            "Add prompt rules that the assistant must not invent temple policies when source data is missing.",
        ],
    )

    add_heading(doc, "12. Step-by-step implementation plan", 1)
    add_numbered(
        doc,
        [
            "Create temple-ai-service as a new FastAPI microservice.",
            "Create a new temple_ai Postgres database and enable pgvector.",
            "Add ingestion jobs for temple notices, wall of fame, payment instructions, FAQs, and booking rules.",
            "Implement chunking and embeddings with text-embedding-3-small.",
            "Implement retrieval API filtered by temple_id and role visibility.",
            "Integrate OpenAI Responses API with GPT-4.1 and structured tool calling.",
            "Build Chat tab UI with message list, assistant responses, and action cards.",
            "Add devotee tools first: membership status, booking status, donation status.",
            "Add admin drafting tools second: draft notification and draft wall of fame.",
            "Add citations, logging, and quality review dashboards.",
            "Run evaluation on real temple questions before wider rollout.",
        ],
    )

    add_heading(doc, "13. Learning roadmap", 1)
    add_paragraph(
        doc,
        "Use this order if the team wants to learn and implement step by step without getting overwhelmed.",
        space_after=8,
    )
    add_numbered(
        doc,
        [
            "Learn embeddings and vector search fundamentals.",
            "Learn the OpenAI Responses API and tool calling.",
            "Build simple retrieval first before any agent behavior.",
            "Add temple-scoped metadata filters.",
            "Add one live tool at a time.",
            "Only then build richer agentic UI actions.",
        ],
    )
    add_two_col_table(
        doc,
        [
            ("OpenAI Responses API", "https://platform.openai.com/docs/api-reference/responses/retrieve"),
            ("OpenAI file search guide", "https://platform.openai.com/docs/guides/tools-file-search/"),
            ("OpenAI GPT-4.1 model page", "https://platform.openai.com/docs/models/gpt-4.1"),
            ("OpenAI embeddings guide", "https://platform.openai.com/docs/guides/embeddings/what-are-embeddings"),
            ("text-embedding-3-small", "https://platform.openai.com/docs/models/text-embedding-3-small"),
            ("text-embedding-3-large", "https://platform.openai.com/docs/models/text-embedding-3-large"),
            ("OpenAI retrieval guide", "https://platform.openai.com/docs/guides/retrieval"),
            ("pgvector project", "https://github.com/pgvector/pgvector"),
        ],
        ["Learning topic", "Reference"],
    )

    add_heading(doc, "14. Terminology glossary", 1)
    add_two_col_table(
        doc,
        [
            ("RAG", "Retrieval-Augmented Generation. The model answers using retrieved project or domain data instead of relying only on pretraining."),
            ("Embedding", "A numeric vector representation of text used for semantic similarity search."),
            ("Vector store", "A database that stores embeddings and supports similarity search."),
            ("Chunking", "Splitting larger content into smaller passages so retrieval can find the most relevant part."),
            ("Metadata filter", "A filter such as temple_id, language, or role that limits which chunks can be retrieved."),
            ("Retrieval", "The search step that finds relevant chunks before answer generation."),
            ("Reranking", "A second pass that reorders retrieved chunks for better relevance."),
            ("Grounding", "Making the answer rely on retrieved or tool-backed facts instead of unsupported guesses."),
            ("Hallucination", "When the model states something confidently that is not supported by source data."),
            ("Tool calling", "Letting the model ask your application to run a defined function or API action."),
            ("Agent", "A model workflow that can reason, choose tools, and continue toward a goal."),
            ("Agentic UI", "A user interface where assistant responses include actions, cards, buttons, or workflow steps, not only plain text."),
            ("Structured output", "A response format with fixed fields that can be safely parsed by the app."),
            ("Prompt", "The instructions and context given to the model."),
            ("Context window", "The maximum amount of input text the model can consider in one request."),
            ("Ingestion", "The pipeline that reads source data and prepares it for retrieval."),
            ("Citation", "A source reference shown with the answer to improve trust and traceability."),
            ("Tenant isolation", "Keeping one temple's data separated from another temple's data."),
        ],
        ["Term", "Definition"],
    )

    add_heading(doc, "15. Final recommendation", 1)
    add_info_box(
        doc,
        "Recommended first release",
        "Build a temple-scoped assistant in a new temple-ai-service using GPT-4.1 for responses and text-embedding-3-small for retrieval. Start with temple knowledge questions plus three live tools: membership status, booking status, and donation status. Add admin drafting help after that.",
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
